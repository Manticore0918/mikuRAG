import re
import zipfile
from dataclasses import replace
from pathlib import Path
from typing import Any, Protocol
from xml.etree import ElementTree

from docx import Document as DocxDocument
from docx.table import Table
from docx.text.paragraph import Paragraph
from pypdf import PdfReader

from app.ingestion.contracts import (
    ExtractedBlock,
    ExtractedDocument,
    ExtractedSection,
    ExtractionWarning,
)
from app.ingestion.errors import ExtractionError

_HEADING_STYLE = re.compile(r"^Heading\s+([1-6])$", re.IGNORECASE)
_MARKDOWN_HEADING = re.compile(r"^(#{1,6})\s+(.+?)\s*$")
_LIST_ITEM = re.compile(r"^\s*(?:[-*+]\s+|\d+[.)]\s+)")
_SECTION_HEADING = re.compile(r"^\s*(?:\d+(?:\.\d+)*[.)]?\s+)?\S.*$")


class PdfFallbackExtractor(Protocol):
    def extract_page(self, path: Path, page_number: int) -> str | None: ...


def _require_content(blocks: list[ExtractedBlock]) -> None:
    if not any(block.text.strip() for block in blocks):
        raise ExtractionError(
            "No extractable text was found. Scanned or image-only Documents require OCR, "
            "which is not supported in this deployment."
        )


def extract_pdf(
    path: Path,
    max_pages: int,
    fallback_extractor: PdfFallbackExtractor | None = None,
) -> ExtractedDocument:
    try:
        reader = PdfReader(path)
        if reader.is_encrypted:
            raise ExtractionError("Password-protected PDF Documents are not supported")
        page_count = len(reader.pages)
        if page_count > max_pages:
            raise ExtractionError(f"Documents cannot exceed {max_pages} pages")

        blocks: list[ExtractedBlock] = []
        warnings: list[ExtractionWarning] = []
        heading_path: list[str] = []
        for page_number, page in enumerate(reader.pages, start=1):
            text = page.extract_text() or ""
            fast_text_is_sparse = len(re.sub(r"\s+", "", text)) < 80
            if fast_text_is_sparse and fallback_extractor is not None:
                fallback_text = fallback_extractor.extract_page(path, page_number) or ""
                if len(fallback_text.strip()) > len(text.strip()):
                    text = fallback_text
                    warnings.append(
                        ExtractionWarning(
                            code="ocr_fallback_used",
                            message="A fallback extractor supplied text for this page.",
                            page_number=page_number,
                        )
                    )
            if not text.strip():
                warnings.append(
                    ExtractionWarning(
                        code="empty_page",
                        message="No text was extracted from this page.",
                        page_number=page_number,
                    )
                )
                continue
            if len(re.sub(r"\s+", "", text)) < 80:
                warnings.append(
                    ExtractionWarning(
                        code="low_text_density",
                        message="The page has unusually little extractable text.",
                        page_number=page_number,
                    )
                )
            if _suspected_multi_column(text):
                warnings.append(
                    ExtractionWarning(
                        code="suspected_multi_column",
                        message="The extracted reading order may contain multiple columns.",
                        page_number=page_number,
                    )
                )
            page_blocks, heading_path = _pdf_page_blocks(
                text,
                page_number=page_number,
                start_order=len(blocks),
                heading_path=heading_path,
            )
            blocks.extend(page_blocks)
    except ExtractionError:
        raise
    except Exception as error:
        raise ExtractionError("The PDF could not be read safely") from error
    _require_content(blocks)
    return ExtractedDocument(
        blocks=blocks,
        page_count=page_count,
        warnings=warnings,
    )


def _suspected_multi_column(text: str) -> bool:
    lines = [line for line in text.splitlines() if line.strip()]
    if len(lines) < 5:
        return False
    split_lines = sum(bool(re.search(r"\S {4,}\S", line)) for line in lines)
    return split_lines / len(lines) >= 0.3


def _pdf_page_blocks(
    text: str,
    *,
    page_number: int,
    start_order: int,
    heading_path: list[str],
) -> tuple[list[ExtractedBlock], list[str]]:
    raw_groups = [
        [line.strip() for line in group.splitlines() if line.strip()]
        for group in re.split(r"\n\s*\n", text)
        if group.strip()
    ]
    if not raw_groups:
        return [], heading_path

    blocks: list[ExtractedBlock] = []
    current_path = list(heading_path)
    paragraph_lines: list[str] = []
    table_lines: list[str] = []

    def append_block(
        block_text: str,
        block_type: str,
        *,
        source_group: int,
        heading_level: int | None = None,
    ) -> None:
        blocks.append(
            ExtractedBlock(
                text=block_text,
                block_type=block_type,
                order=start_order + len(blocks),
                start_page=page_number,
                end_page=page_number,
                heading_level=heading_level,
                heading_path=list(current_path),
                metadata={
                    "locator": {"page": page_number},
                    "source": "pypdf",
                    "source_group": source_group,
                },
            )
        )

    def flush_paragraph(source_group: int) -> None:
        nonlocal paragraph_lines
        if paragraph_lines:
            append_block(
                "\n".join(paragraph_lines),
                "paragraph",
                source_group=source_group,
            )
            paragraph_lines = []

    def flush_table(source_group: int) -> None:
        nonlocal table_lines
        if table_lines:
            append_block(
                "\n".join(table_lines),
                "table",
                source_group=source_group,
            )
            table_lines = []

    for group_number, lines in enumerate(raw_groups):
        for line in lines:
            block_type = _line_block_type(line)
            if block_type == "paragraph":
                flush_table(group_number)
                paragraph_lines.append(line)
                continue
            flush_paragraph(group_number)
            if block_type == "table":
                table_lines.append(line)
                continue
            flush_table(group_number)
            if block_type == "heading":
                heading_level = _inferred_heading_level(line)
                current_path = _updated_heading_path(current_path, heading_level, line)
                append_block(
                    line,
                    block_type,
                    source_group=group_number,
                    heading_level=heading_level,
                )
            else:
                append_block(
                    line,
                    block_type,
                    source_group=group_number,
                )
        flush_paragraph(group_number)
        flush_table(group_number)

    if blocks and len(blocks[0].text) <= 160:
        metadata = dict(blocks[0].metadata)
        metadata["page_position"] = "header"
        blocks[0] = replace(blocks[0], metadata=metadata)
    if blocks and len(blocks[-1].text) <= 160:
        metadata = dict(blocks[-1].metadata)
        metadata["page_position"] = "footer"
        blocks[-1] = replace(blocks[-1], metadata=metadata)
    return blocks, current_path


def _line_block_type(line: str) -> str:
    if _LIST_ITEM.match(line):
        return "list_item"
    if "|" in line or "\t" in line:
        return "table"
    if _looks_like_heading(line):
        return "heading"
    return "paragraph"


def _looks_like_heading(line: str) -> bool:
    stripped = line.strip()
    words = stripped.split()
    if not stripped or len(stripped) > 120 or len(words) > 14:
        return False
    if stripped.endswith((".", ",", ";", "?", "!")):
        return False
    if not _SECTION_HEADING.match(stripped):
        return False
    letters = [character for character in stripped if character.isalpha()]
    if letters and all(character.isupper() for character in letters):
        return True
    return stripped.istitle() or bool(re.match(r"^\d+(?:\.\d+)+\s+\S", stripped))


def _inferred_heading_level(line: str) -> int:
    numbered = re.match(r"^(\d+(?:\.\d+)*)", line.strip())
    if numbered:
        return min(numbered.group(1).count(".") + 1, 6)
    return 1


def _updated_heading_path(path: list[str], level: int, heading: str) -> list[str]:
    prefix = path[: level - 1]
    return [*prefix, heading.strip()]


def _docx_page_count(path: Path) -> int | None:
    try:
        with zipfile.ZipFile(path) as archive:
            root = ElementTree.fromstring(archive.read("docProps/app.xml"))
        pages = root.find(
            "{http://schemas.openxmlformats.org/officeDocument/2006/extended-properties}Pages"
        )
        return int(pages.text) if pages is not None and pages.text else None
    except (KeyError, ValueError, zipfile.BadZipFile, ElementTree.ParseError):
        return None


def extract_docx(path: Path, max_pages: int) -> ExtractedDocument:
    page_count = _docx_page_count(path)
    if page_count is not None and page_count > max_pages:
        raise ExtractionError(f"Documents cannot exceed {max_pages} pages")
    try:
        document = DocxDocument(path)
        blocks: list[ExtractedBlock] = []
        heading_path: list[str] = []
        for source_order, block in enumerate(document.iter_inner_content()):
            heading_level = None
            if isinstance(block, Paragraph):
                text = block.text.strip()
                style_name = block.style.name if block.style is not None else ""
                heading_match = _HEADING_STYLE.match(style_name)
                if heading_match:
                    block_type = "heading"
                    heading_level = int(heading_match.group(1))
                    heading_path = _updated_heading_path(
                        heading_path,
                        heading_level,
                        text,
                    )
                elif "list" in style_name.casefold() or _LIST_ITEM.match(text):
                    block_type = "list_item"
                elif "code" in style_name.casefold() or "preformatted" in style_name.casefold():
                    block_type = "code"
                else:
                    block_type = "paragraph"
                metadata: dict[str, Any] = {"style": style_name}
            elif isinstance(block, Table):
                text = "\n".join(
                    " | ".join(cell.text.strip() for cell in row.cells) for row in block.rows
                ).strip()
                block_type = "table"
                metadata = {
                    "rows": len(block.rows),
                    "columns": max((len(row.cells) for row in block.rows), default=0),
                }
            else:
                continue
            if text:
                section_number = len(blocks) + 1
                metadata["source_order"] = source_order
                metadata["locator"] = {"section": section_number}
                blocks.append(
                    ExtractedBlock(
                        text=text,
                        block_type=block_type,
                        order=len(blocks),
                        heading_level=heading_level,
                        heading_path=list(heading_path),
                        metadata=metadata,
                    )
                )
    except Exception as error:
        raise ExtractionError("The DOCX Document could not be read safely") from error
    _require_content(blocks)
    return ExtractedDocument(blocks=blocks, page_count=page_count)


def extract_text(path: Path, markdown: bool) -> ExtractedDocument:
    try:
        content = path.read_text(encoding="utf-8-sig")
    except (OSError, UnicodeDecodeError) as error:
        raise ExtractionError("The text Document could not be read as UTF-8") from error

    blocks = _extract_markdown_blocks(content) if markdown else _extract_plain_text_blocks(content)
    _require_content(blocks)
    return ExtractedDocument(blocks=blocks, page_count=None)


def _extract_plain_text_blocks(content: str) -> list[ExtractedBlock]:
    blocks: list[ExtractedBlock] = []
    matches = re.finditer(r"\S(?:.*?\S)?(?=\n\s*\n|\s*\Z)", content, re.S)
    for match in matches:
        line_start = content.count("\n", 0, match.start()) + 1
        line_end = content.count("\n", 0, match.end()) + 1
        section_number = len(blocks) + 1
        blocks.append(
            ExtractedBlock(
                text=match.group(0),
                block_type="paragraph",
                order=len(blocks),
                metadata={
                    "locator": {
                        "section": section_number,
                        "line_start": line_start,
                        "line_end": line_end,
                    }
                },
            )
        )
    return blocks


def _extract_markdown_blocks(content: str) -> list[ExtractedBlock]:
    blocks: list[ExtractedBlock] = []
    heading_path: list[str] = []
    current_section: str | int = "Document"
    paragraph_lines: list[str] = []
    paragraph_start = 1
    table_lines: list[str] = []
    table_start = 1
    code_lines: list[str] = []
    code_start = 1
    in_code = False

    def append_block(
        text: str,
        block_type: str,
        line_start: int,
        line_end: int,
        heading_level: int | None = None,
    ) -> None:
        if not text.strip():
            return
        blocks.append(
            ExtractedBlock(
                text=text,
                block_type=block_type,
                order=len(blocks),
                heading_level=heading_level,
                heading_path=list(heading_path),
                metadata={
                    "locator": {
                        "section": current_section,
                        "line_start": line_start,
                        "line_end": line_end,
                    }
                },
            )
        )

    def flush_paragraph(end_line: int) -> None:
        nonlocal paragraph_lines
        append_block("\n".join(paragraph_lines), "paragraph", paragraph_start, end_line)
        paragraph_lines = []

    def flush_table(end_line: int) -> None:
        nonlocal table_lines
        append_block("\n".join(table_lines), "table", table_start, end_line)
        table_lines = []

    lines = content.splitlines()
    for line_number, line in enumerate(lines, start=1):
        if line.lstrip().startswith("```"):
            if in_code:
                append_block("\n".join(code_lines), "code", code_start, line_number)
                code_lines = []
                in_code = False
            else:
                flush_paragraph(line_number - 1)
                flush_table(line_number - 1)
                in_code = True
                code_start = line_number
            continue
        if in_code:
            code_lines.append(line)
            continue

        heading_match = _MARKDOWN_HEADING.match(line)
        if heading_match:
            flush_paragraph(line_number - 1)
            flush_table(line_number - 1)
            level = len(heading_match.group(1))
            heading = heading_match.group(2).strip()
            heading_path = _updated_heading_path(heading_path, level, heading)
            current_section = heading
            append_block(heading, "heading", line_number, line_number, level)
            continue
        if not line.strip():
            flush_paragraph(line_number - 1)
            flush_table(line_number - 1)
            continue
        if _LIST_ITEM.match(line):
            flush_paragraph(line_number - 1)
            flush_table(line_number - 1)
            append_block(line.strip(), "list_item", line_number, line_number)
            continue
        if "|" in line and line.count("|") >= 2:
            flush_paragraph(line_number - 1)
            if not table_lines:
                table_start = line_number
            table_lines.append(line.strip())
            continue
        flush_table(line_number - 1)
        if not paragraph_lines:
            paragraph_start = line_number
        paragraph_lines.append(line)

    if in_code:
        append_block("\n".join(code_lines), "code", code_start, max(1, len(lines)))
    else:
        flush_paragraph(max(1, len(lines)))
        flush_table(max(1, len(lines)))
    return blocks


def extract_document(path: Path, media_type: str, max_pages: int) -> ExtractedDocument:
    if media_type == "application/pdf":
        return extract_pdf(path, max_pages)
    if media_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        return extract_docx(path, max_pages)
    return extract_text(path, markdown=media_type == "text/markdown")


__all__ = [
    "ExtractedBlock",
    "ExtractedDocument",
    "ExtractedSection",
    "ExtractionWarning",
    "PdfFallbackExtractor",
    "extract_document",
    "extract_docx",
    "extract_pdf",
    "extract_text",
]
