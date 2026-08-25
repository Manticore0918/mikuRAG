from pathlib import Path

import pytest
from docx import Document as DocxDocument
from pypdf import PdfWriter

from app.ingestion.chunking import chunk_sections
from app.ingestion.errors import ExtractionError
from app.ingestion.extraction import EXTRACTOR_REGISTRY, ExtractedSection, extract_document


def test_plain_text_extraction_accepts_a_single_trailing_newline(tmp_path: Path) -> None:
    path = tmp_path / "restart-proof.txt"
    path.write_text("The upload checkpoint survived.\n", encoding="utf-8")

    extracted = extract_document(path, "text/plain", 500)

    assert [block.text for block in extracted.blocks] == [
        "The upload checkpoint survived."
    ]
    assert extracted.parser_version == "plain_text_v1"


def test_extractor_registry_covers_every_checkpoint_one_media_type() -> None:
    assert {
        "application/pdf",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "text/plain",
        "text/markdown",
        "text/html",
        "text/x-python",
        "text/javascript",
        "text/typescript",
        "text/jsx",
        "text/tsx",
    }.issubset(EXTRACTOR_REGISTRY.media_types)


def test_html_extraction_removes_noise_and_keeps_title_heading_and_element(
    tmp_path: Path,
) -> None:
    path = tmp_path / "runbook.html"
    path.write_text(
        """<!doctype html>
<html><head><title>Recovery Runbook</title>
<link rel="canonical" href="https://docs.example.test/recovery"></head>
<body><nav>Ignore this menu</nav><main id="content"><h1>Restore</h1>
<p>Use backup MIKU-4271.</p><script>steal('secret')</script></main></body></html>""",
        encoding="utf-8",
    )

    extracted = extract_document(path, "text/html", 500, source_kind="html")

    combined = " ".join(block.text for block in extracted.blocks)
    assert "Recovery Runbook" in combined
    assert "Use backup MIKU-4271" in combined
    assert "Ignore this menu" not in combined
    assert "steal" not in combined
    assert extracted.parser_version == "html_stdlib_v1"
    assert extracted.metadata == {
        "title": "Recovery Runbook",
        "canonical_uri": "https://docs.example.test/recovery",
    }
    paragraph = next(block for block in extracted.blocks if "MIKU-4271" in block.text)
    assert paragraph.heading_path == ["Restore"]
    assert "#content" in paragraph.metadata["locator"]["element"]
    assert paragraph.metadata["locator"]["line_start"] == 5
    assert paragraph.metadata["locator"]["text_start"] > 0
    assert [warning.code for warning in extracted.warnings] == ["html_navigation_removed"]


def test_python_extraction_keeps_repo_path_module_symbols_and_line_ranges(
    tmp_path: Path,
) -> None:
    path = tmp_path / "worker.py"
    path.write_text(
        "import asyncio\n\n"
        "class Worker:\n"
        "    def run(self):\n"
        "        return 'MIKU-4271'\n\n"
        "def healthcheck():\n"
        "    return True\n",
        encoding="utf-8",
    )

    extracted = extract_document(
        path,
        "text/x-python",
        500,
        source_kind="code",
        source_path="src/jobs/worker.py",
        language="python",
    )

    worker = next(
        block
        for block in extracted.blocks
        if block.metadata["locator"].get("symbol") == "Worker"
    )
    locator = worker.metadata["locator"]
    assert locator["path"] == "src/jobs/worker.py"
    assert locator["module"] == "src.jobs.worker"
    assert locator["line_start"] == 3
    assert locator["line_end"] == 5
    assert extracted.parser_version == "python_ast_v1"


def test_malformed_python_has_a_safe_parser_error(tmp_path: Path) -> None:
    path = tmp_path / "broken.py"
    path.write_text("def broken(:\n    pass\n", encoding="utf-8")

    with pytest.raises(ExtractionError, match="could not be parsed safely near line 1"):
        extract_document(path, "text/x-python", 500, source_kind="code")


def test_typescript_extraction_detects_top_level_symbols(tmp_path: Path) -> None:
    path = tmp_path / "client.ts"
    path.write_text(
        "export interface Client { ready: boolean }\n\n"
        "export async function connect() {\n  return 'MIKU-4271'\n}\n",
        encoding="utf-8",
    )

    extracted = extract_document(
        path,
        "text/typescript",
        500,
        source_kind="code",
        source_path="web/client.ts",
        language="typescript",
    )

    assert [block.metadata["locator"].get("symbol") for block in extracted.blocks] == [
        "Client",
        "connect",
    ]
    assert extracted.blocks[1].metadata["locator"]["line_start"] == 3


def test_markdown_extraction_preserves_heading_and_line_locator(tmp_path: Path) -> None:
    path = tmp_path / "guide.md"
    path.write_text(
        "# Setup\n\nInstall the service.\n\n## Verify\nRun the checks.", encoding="utf-8"
    )

    extracted = extract_document(path, "text/markdown", 500)

    assert [section.locator["section"] for section in extracted.sections] == ["Setup", "Verify"]
    assert extracted.sections[0].locator["line_start"] == 1
    assert [block.block_type for block in extracted.blocks] == [
        "heading",
        "paragraph",
        "heading",
        "paragraph",
    ]
    assert extracted.blocks[-1].heading_path == ["Setup", "Verify"]


def test_markdown_extraction_marks_lists_tables_and_code(tmp_path: Path) -> None:
    path = tmp_path / "structured.md"
    path.write_text(
        "# Data\n"
        "- one\n"
        "- two\n\n"
        "| Key | Value |\n"
        "| --- | --- |\n"
        "| A | B |\n\n"
        "```\n"
        "x = 1\n"
        "```\n",
        encoding="utf-8",
    )

    extracted = extract_document(path, "text/markdown", 500)

    assert [block.block_type for block in extracted.blocks] == [
        "heading",
        "list_item",
        "list_item",
        "table",
        "code",
    ]
    assert extracted.blocks[3].text.count("\n") == 2
    assert extracted.blocks[4].text == "x = 1"


def test_docx_extraction_keeps_paragraphs_and_tables(tmp_path: Path) -> None:
    path = tmp_path / "guide.docx"
    document = DocxDocument()
    document.add_heading("Operations", level=1)
    document.add_paragraph("Restart the worker after configuration changes.")
    table = document.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "State"
    table.rows[0].cells[1].text = "Ready"
    document.save(path)

    extracted = extract_document(
        path,
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        500,
    )

    assert any("Restart the worker" in section.text for section in extracted.sections)
    assert any("State | Ready" in section.text for section in extracted.sections)
    assert [block.block_type for block in extracted.blocks] == [
        "heading",
        "paragraph",
        "table",
    ]
    assert extracted.blocks[1].heading_path == ["Operations"]


def test_image_only_pdf_is_rejected_without_partial_content(tmp_path: Path) -> None:
    path = tmp_path / "scan.pdf"
    writer = PdfWriter()
    writer.add_blank_page(width=100, height=100)
    with path.open("wb") as output:
        writer.write(output)

    with pytest.raises(ExtractionError, match="No extractable text"):
        extract_document(path, "application/pdf", 500)


def test_pdf_page_limit_is_enforced_before_extraction(tmp_path: Path) -> None:
    path = tmp_path / "large.pdf"
    writer = PdfWriter()
    writer.add_blank_page(width=100, height=100)
    writer.add_blank_page(width=100, height=100)
    with path.open("wb") as output:
        writer.write(output)

    with pytest.raises(ExtractionError, match="cannot exceed 1 pages"):
        extract_document(path, "application/pdf", 1)


def test_chunking_never_crosses_section_boundaries_and_retains_overlap() -> None:
    sections = [
        ExtractedSection(text="alpha " * 80, locator={"page": 1}),
        ExtractedSection(text="beta " * 20, locator={"page": 2}),
    ]

    chunks = chunk_sections(sections, target_characters=120, overlap_characters=20)

    assert len(chunks) > 2
    assert all(chunk.locator["page"] in {1, 2} for chunk in chunks)
    assert not any("alpha" in chunk.text and "beta" in chunk.text for chunk in chunks)
    assert [chunk.locator["part"] for chunk in chunks if chunk.locator["page"] == 2] == [1]
