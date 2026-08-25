import uuid
from pathlib import Path

import pytest
from docx import Document as DocxDocument
from sqlalchemy.sql.dml import Delete

import app.ingestion.extraction as extraction
from app.ingestion.chunking import TextChunk
from app.ingestion.contracts import ExtractedDocument
from app.ingestion.extraction import extract_document, extract_pdf
from app.ingestion.hierarchical_chunking import (
    HierarchicalChunkingConfig,
    construct_hierarchy,
)
from app.ingestion.normalization import normalize_document
from app.ingestion.persistence import (
    build_hierarchical_chunk_models,
    build_legacy_chunk_models,
    replace_document_chunks,
)
from app.ingestion.tokenization import ConservativeTokenizer
from app.ingestion.validation import validate_hierarchy


class _PdfPage:
    def __init__(self, text: str) -> None:
        self._text = text

    def extract_text(self) -> str:
        return self._text


class _PdfReader:
    is_encrypted = False

    def __init__(self, texts: list[str]) -> None:
        self.pages = [_PdfPage(text) for text in texts]


def _config() -> HierarchicalChunkingConfig:
    return HierarchicalChunkingConfig(
        child_min_tokens=4,
        child_target_tokens=12,
        child_max_tokens=20,
        child_overlap_tokens=2,
        parent_target_tokens=40,
        parent_max_tokens=64,
        chunking_version="hierarchical_v1",
    )


def _run_hierarchy(document: ExtractedDocument):
    tokenizer = ConservativeTokenizer()
    normalized = normalize_document(document)
    hierarchy = construct_hierarchy(
        normalized,
        config=_config(),
        tokenizer=tokenizer,
    )
    validate_hierarchy(
        normalized,
        hierarchy,
        config=_config(),
        tokenizer=tokenizer,
        max_document_chunks=20_000,
        max_document_tokens=500_000,
    )
    return normalized, hierarchy


def test_simple_and_200_page_pdfs_complete_the_chunking_pipeline(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    simple_pages = [
        "Operations Guide\n\n"
        "Restart the worker after every configuration change and verify its health."
    ]
    monkeypatch.setattr(extraction, "PdfReader", lambda _: _PdfReader(simple_pages))
    simple, simple_hierarchy = _run_hierarchy(
        extract_pdf(Path("simple.pdf"), max_pages=500)
    )

    large_pages = [
        f"Page {page}\n\n"
        f"Marker{page} describes the operational requirement for this part of the handbook "
        "and provides enough source content for deterministic extraction."
        for page in range(1, 201)
    ]
    monkeypatch.setattr(extraction, "PdfReader", lambda _: _PdfReader(large_pages))
    large, large_hierarchy = _run_hierarchy(
        extract_pdf(Path("large.pdf"), max_pages=200)
    )

    assert simple.page_count == 1
    assert simple_hierarchy.children
    assert large.page_count == 200
    assert {block.start_page for block in large.blocks} == set(range(1, 201))
    covered_orders = {
        order
        for parent in large_hierarchy.parents
        for order in parent.source_block_orders
    }
    assert covered_orders == {
        block.order for block in large.blocks
    }


def test_cross_page_multicolumn_and_spanning_table_pdf_cases(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    cross_page = [
        "The policy describes the inter-",
        "national approval requirement and the controls that apply.",
    ]
    monkeypatch.setattr(extraction, "PdfReader", lambda _: _PdfReader(cross_page))
    cross_normalized, cross_hierarchy = _run_hierarchy(
        extract_pdf(Path("cross-page.pdf"), max_pages=10)
    )

    multi_column = [
        "\n".join(
            [
                "Left one    Right one",
                "Left two    Right two",
                "Left three    Right three",
                "Left four    Right four",
                "Left five    Right five",
            ]
        )
    ]
    monkeypatch.setattr(extraction, "PdfReader", lambda _: _PdfReader(multi_column))
    columns = extract_pdf(Path("columns.pdf"), max_pages=10)

    table_pages = [
        "Key | Value\nA | 1\nB | 2",
        "Key | Value\nC | 3\nD | 4",
    ]
    monkeypatch.setattr(extraction, "PdfReader", lambda _: _PdfReader(table_pages))
    table_normalized, table_hierarchy = _run_hierarchy(
        extract_pdf(Path("table.pdf"), max_pages=10)
    )

    assert cross_normalized.blocks[0].text.startswith("The policy describes the international")
    assert (cross_normalized.blocks[0].start_page, cross_normalized.blocks[0].end_page) == (
        1,
        2,
    )
    assert any((child.start_page, child.end_page) == (1, 2) for child in cross_hierarchy.children)
    assert "suspected_multi_column" in {warning.code for warning in columns.warnings}
    assert [block.block_type for block in table_normalized.blocks] == ["table", "table"]
    table_pages = {
        page
        for child in table_hierarchy.children
        for page in (child.start_page, child.end_page)
    }
    assert table_pages == {
        1,
        2,
    }


def test_scanned_pdf_uses_ocr_fallback_before_chunking(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    class Fallback:
        def extract_page(self, path: Path, page_number: int) -> str:
            return (
                "Recovered scanner text contains the approval procedure and enough detail "
                "to be treated as searchable source evidence."
            )

    monkeypatch.setattr(extraction, "PdfReader", lambda _: _PdfReader([""]))

    extracted = extract_pdf(
        Path("scan.pdf"),
        max_pages=10,
        fallback_extractor=Fallback(),
    )
    normalized, hierarchy = _run_hierarchy(extracted)

    assert [warning.code for warning in normalized.warnings] == ["ocr_fallback_used"]
    assert hierarchy.children[0].text.startswith("Recovered scanner text")


def test_mixed_pdf_docx_markdown_and_text_knowledge_base(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    text_path = tmp_path / "notes.txt"
    text_path.write_text("Plain text operational notes.", encoding="utf-8")
    markdown_path = tmp_path / "guide.md"
    markdown_path.write_text("# Setup\n\nInstall and verify the service.", encoding="utf-8")
    docx_path = tmp_path / "handbook.docx"
    docx = DocxDocument()
    docx.add_heading("Recovery", level=1)
    docx.add_paragraph("Restore the latest verified backup.")
    docx.save(docx_path)
    monkeypatch.setattr(
        extraction,
        "PdfReader",
        lambda _: _PdfReader(
            ["PDF policy text explains the escalation process in sufficient detail."]
        ),
    )

    documents = [
        extract_document(Path("policy.pdf"), "application/pdf", 10),
        extract_document(
            docx_path,
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            10,
        ),
        extract_document(markdown_path, "text/markdown", 10),
        extract_document(text_path, "text/plain", 10),
    ]
    outputs = [_run_hierarchy(document)[1] for document in documents]

    assert len(outputs) == 4
    assert all(output.children for output in outputs)
    assert any(not child.heading_path for child in outputs[-1].children)


@pytest.mark.asyncio
async def test_reingestion_replaces_legacy_chunks_with_hierarchical_models() -> None:
    document_id = uuid.uuid4()
    tokenizer = ConservativeTokenizer()
    legacy = build_legacy_chunk_models(
        document_id=document_id,
        chunks=[TextChunk(text="Legacy source text.", locator={"page": 1})],
        vectors=[[0.1] * 768],
        embedding_model="embed-v1",
        tokenizer=tokenizer,
    )
    normalized, hierarchy = _run_hierarchy(
        ExtractedDocument(
            blocks=[
                extraction.ExtractedBlock(
                    text="Replacement source text with hierarchical context.",
                    block_type="paragraph",
                    order=0,
                    start_page=1,
                    end_page=1,
                )
            ],
            page_count=1,
        )
    )
    hierarchical = build_hierarchical_chunk_models(
        document_id=document_id,
        hierarchy=hierarchy,
        vectors=[[0.2] * 768 for _ in hierarchy.children],
        embedding_model="embed-v1",
    )

    class Session:
        def __init__(self) -> None:
            self.statements = []
            self.batches = []

        async def execute(self, statement) -> None:
            self.statements.append(statement)

        def add_all(self, values) -> None:
            self.batches.append(list(values))

        async def flush(self) -> None:
            pass

    session = Session()
    await replace_document_chunks(session, document_id=document_id, batch=legacy)
    await replace_document_chunks(session, document_id=document_id, batch=hierarchical)

    assert normalized.blocks[0].text.startswith("Replacement source")
    assert sum(isinstance(statement, Delete) for statement in session.statements) == 2
    assert session.batches == [
        legacy.children,
        hierarchical.parents,
        hierarchical.children,
    ]
    assert {chunk.chunking_version for chunk in legacy.children} == {"legacy"}
    assert {chunk.chunking_version for chunk in hierarchical.children} == {
        "hierarchical_v1"
    }
    assert not {chunk.id for chunk in legacy.children} & {
        chunk.id for chunk in hierarchical.children
    }
